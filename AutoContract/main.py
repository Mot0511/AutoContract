from flask import *
from docx import Document
from docx.shared import Inches

app = Flask(__name__, static_folder="static")
# -*- coding: utf8 -*-
document = Document()

# @app.route('/ok', methods=['GET', 'POST'])
# def test():
#     if request.method == 'POST':
#         name = request.form.get('name')
#     return render_template('ok.html', name=name)

@app.route('/')
def index(name=None):
    return render_template('index.html')

@app.route('/contract', methods=['GET', 'POST'])
def contract():
    if request.method == 'POST':
        place = request.form.get('place')
        date1_sost = request.form.get('date1_sost')
        name_saler = request.form.get('name_saler')
        date_birth_saler = request.form.get('date_birth_saler')
        place_birth_saler = request.form.get('place_birth_saler')
        date_pass_saler = request.form.get('date_pass_saler')
        number_pass_saler = request.form.get('number_pass_saler')
        seria_saler = request.form.get('seria_saler')
        vadan_saler = request.form.get('vadan_saler')
        adres_saler = request.form.get('adres_saler')

        name_bayer = request.form.get('name_bayer')
        tel_bayer = request.form.get('tel_bayer')
        date_birth_bayer = request.form.get('date_birth_bayer')
        place_birth_bayer = request.form.get('place_birth_bayer')
        date_pass_bayer = request.form.get('date_pass_bayer')
        number_pass_bayer = request.form.get('number_pass_bayer')
        seria_bayer = request.form.get('seria_bayer')
        vadan_bayer = request.form.get('vadan_bayer')
        adress_bayer = request.form.get('adress_bayer')

        mm_car = request.form.get('mm_car')
        vin = request.form.get('vin')
        type_tc = request.form.get('type_tc')
        model_dvig = request.form.get('model_dvig')
        number_dvig = request.form.get('number_dvig')
        number_shassi = request.form.get('number_shassi')
        number_kusov = request.form.get('number_kusov')
        color = request.form.get('color')
        power = request.form.get('power')
        volume = request.form.get('volume')
        date_tc = request.form.get('date_tc')
        seria_tc = request.form.get('seria_tc')
        number_tc = request.form.get('number_tc')
        vadan_tc = request.form.get('vadan_tc')
        date_ctc = request.form.get('date_ctc')
        seria_tc = request.form.get('seria_tc')
        number_ctc = request.form.get('number_ctc')
        vadan_ctc = request.form.get('vadan_ctc')
        probeg = request.form.get('probeg')
        gosnumber = request.form.get('gosnumber')
        price = request.form.get('price')
        seria_ctc = request.form.get('seria_ctc')


        document.add_paragraph(place + '/' + date1_sost)

        document.add_heading('Договор о купли-продажи авто', 0)
        p = document.add_paragraph(name_bayer + ', ' + date_birth_bayer + ', ' + place_birth_bayer + ', ' + adress_bayer + ', ' + seria_bayer + ', ' + number_pass_bayer + ', ' + vadan_bayer + ', ' + date_birth_bayer + ', ' + date_pass_bayer)

        p2 = document.add_paragraph('именуемый(ая) в дальнейшем «Покупатель» с одной стороны, и')



        p3 = document.add_paragraph(name_saler + ', ' + date_birth_saler + ', ' + place_birth_saler + ', ' + adres_saler + ', ' + seria_saler + ', ' + number_pass_saler + ', ' + vadan_saler + ', ' + date_birth_saler)

        p4 = document.add_paragraph('именуемый(ая) в дальнейшем «Продавец», с другой стороны, именуемые далее при совместном упоминании «Стороны», а по отдельности «Сторона», заключили настоящий договор (далее – «Договор») о нижеследующем')

        document.add_heading('1. Предмет договора')
        p5 = document.add_paragraph('1.1 Продавец обязуется передать в собственность Покупателя, а Покупатель – принять и оплатить транспортное средство (далее – ТС).')

        document.add_paragraph(mm_car)
        document.add_paragraph(vin)
        document.add_paragraph(type_tc)
        document.add_paragraph(model_dvig + '/' + number_dvig)
        document.add_paragraph(number_shassi)
        document.add_paragraph(number_kusov)
        document.add_paragraph(volume + '/' + power)
        document.add_paragraph(color)
        document.add_paragraph(probeg)
        document.add_paragraph(seria_tc + '/' + number_tc)
        document.add_paragraph(vadan_tc)
        document.add_paragraph(date_tc)
        document.add_paragraph(gosnumber)

        document.add_paragraph('1.2. Собственником ТС до его передачи Покупателю является Продавец (свидетельство о регистрации транспортного средства серия')
        document.add_paragraph(seria_ctc + '№' + number_ctc + 'выдано' + ' ' + vadan_ctc + ' ' + date_ctc)
        document.add_paragraph('Право собственности на ТС переходит к Покупателю с момента подписания настоящего Договора.')
        document.add_paragraph('1.3. Передача ТС осуществляется Продавцом в момент передачи Покупателем Продавцу денежных средств в счет оплаты стоимости ТС согласно п. 2. Договора.')
        document.add_heading('2. Стоимость ТС и порядок расчетов')

        document.add_paragraph('Стоимость ТС составляет ' + price + ' рублей')
        document.add_paragraph('(НДС не облагается). Оплата стоимости ТС производится путем 100% предоплаты (наличным или безналичным расчетом).')

        document.add_heading('3. Гарантии и ответственность')
        document.add_paragraph('Продавец гарантирует Покупателю что:')
        document.add_paragraph('3.1. Продавец является собственником ТС.')
        document.add_paragraph('ТС не является предметом обязательств Продавца перед третьими лицами, в том числе не является предметом залога, в отношении ТС не наложен запрет на совершение регистрационных действий, ТС не находится под арестом, не числится в базах данных МВД России как угнанное или похищенное транспортное средство и не имеет иных обременений.')
        document.add_paragraph('3.3. В случае нарушения гарантий, указанных в п. 3.1 – 3.2 настоящего договора, Продавец обязуется незамедлительно возвратить Покупателю стоимость ТС в полном объеме со дня обнаружения соответствующего нарушения.')

        document.add_heading('4. Заключительные положения')
        document.add_paragraph('Настоящий Договор вступает в силу после его подписания Сторонами и действует до момента полного исполнения Сторонами своих обязательств по Договору. Договор составлен в трех экземплярах, имеющих равную юридическую силу.')

        document.add_heading('5. Подписи сторон')
        document.add_paragraph('__________/____________')
        document.add_paragraph('Денежные средства в сумме ' + price + 'руб. получил')
        document.add_paragraph('__________/____________')

        document.add_paragraph('__________/____________')
        document.add_paragraph('ТС получил  ')
        document.add_paragraph('__________/____________')

        document.add_page_break()
        document.save('test.docx')

    return render_template('ok.html')
app.run()